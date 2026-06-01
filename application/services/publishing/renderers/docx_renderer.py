import logging
import os
import re

logger = logging.getLogger(__name__)

class DocxRenderer:
    def render(self, md_content: str, output_path: str) -> str:
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        except ImportError:
            logger.error("python-docx is not installed. DOCX generation failed.")
            return ""

        doc = Document()
        
        # Simple Markdown parsing for DOCX
        lines = md_content.split('\n')
        
        in_code_block = False
        
        for line in lines:
            if line.startswith('```'):
                in_code_block = not in_code_block
                continue
                
            if in_code_block:
                # Add code as paragraph with different style
                p = doc.add_paragraph(line)
                p.style = 'Normal'
                # Attempt to set monospace font if available, or just keep it simple
                continue
                
            if line.startswith('# '):
                heading = doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('- **') or line.startswith('* **'):
                # Handle bold list items simply
                clean_line = line[2:]
                p = doc.add_paragraph(style='List Bullet')
                
                # Extract bold parts using regex
                parts = re.split(r'(\*\*.*?\*\*)', clean_line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        p.add_run(part)
            elif line.strip() == '---':
                doc.add_paragraph('________________________________________________________')
            elif line.strip() != '':
                # Handle inline bolding
                p = doc.add_paragraph()
                parts = re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        p.add_run(part)

        try:
            doc.save(output_path)
            return output_path
        except Exception as e:
            logger.error(f"Error saving DOCX: {e}")
            return ""
